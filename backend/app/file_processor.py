import io
import PyPDF2
import docx
import openpyxl
import mammoth
from typing import Union

class FileProcessor:
    """
    各種ファイル形式からテキストを抽出するクラス
    """
    
    def extract_text(self, file_content: bytes, file_extension: str) -> str:
        """
        ファイル形式に応じてテキストを抽出
        
        Args:
            file_content: ファイルのバイトデータ
            file_extension: ファイル拡張子 (.pdf, .docx, .doc, .xlsx, .xls)
            
        Returns:
            抽出されたテキスト
        """
        try:
            if file_extension.lower() == '.pdf':
                return self._extract_from_pdf(file_content)
            elif file_extension.lower() in ['.docx']:
                return self._extract_from_docx(file_content)
            elif file_extension.lower() in ['.doc']:
                return self._extract_from_doc(file_content)
            elif file_extension.lower() in ['.xlsx', '.xls']:
                return self._extract_from_excel(file_content)
            else:
                raise ValueError(f"Unsupported file extension: {file_extension}")
        except Exception as e:
            raise Exception(f"Error extracting text from {file_extension} file: {str(e)}")

    def _extract_from_pdf(self, file_content: bytes) -> str:
        """PDFファイルからテキストを抽出"""
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            text = ""
            
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"
            
            return text.strip()
        except Exception as e:
            raise Exception(f"PDF reading error: {str(e)}")

    def _extract_from_docx(self, file_content: bytes) -> str:
        """DOCXファイルからテキストを抽出"""
        try:
            doc = docx.Document(io.BytesIO(file_content))
            text = ""
            
            # 段落のテキストを抽出
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # テーブルのテキストも抽出
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\t"
                    text += "\n"
            
            return text.strip()
        except Exception as e:
            raise Exception(f"DOCX reading error: {str(e)}")

    def _extract_from_doc(self, file_content: bytes) -> str:
        """DOCファイルからテキストを抽出（mammothを使用）"""
        try:
            result = mammoth.extract_raw_text(io.BytesIO(file_content))
            return result.value.strip()
        except Exception as e:
            raise Exception(f"DOC reading error: {str(e)}")

    def _extract_from_excel(self, file_content: bytes) -> str:
        """Excelファイルからテキストを抽出"""
        try:
            workbook = openpyxl.load_workbook(io.BytesIO(file_content), data_only=True)
            text = ""
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text += f"シート名: {sheet_name}\n"
                
                for row in sheet.iter_rows(values_only=True):
                    row_text = []
                    for cell in row:
                        if cell is not None:
                            row_text.append(str(cell))
                    if row_text:
                        text += "\t".join(row_text) + "\n"
                text += "\n"
            
            return text.strip()
        except Exception as e:
            raise Exception(f"Excel reading error: {str(e)}")

    def validate_file_size(self, file_content: bytes, max_size_mb: int = 10) -> bool:
        """
        ファイルサイズの検証
        
        Args:
            file_content: ファイルのバイトデータ
            max_size_mb: 最大サイズ（MB）
            
        Returns:
            サイズが許可範囲内かどうか
        """
        file_size_mb = len(file_content) / (1024 * 1024)
        return file_size_mb <= max_size_mb

    def get_file_info(self, file_content: bytes, filename: str) -> dict:
        """
        ファイル情報を取得
        
        Args:
            file_content: ファイルのバイトデータ
            filename: ファイル名
            
        Returns:
            ファイル情報の辞書
        """
        file_size_mb = len(file_content) / (1024 * 1024)
        file_extension = filename.split('.')[-1].lower() if '.' in filename else ''
        
        return {
            "filename": filename,
            "file_size_mb": round(file_size_mb, 2),
            "file_extension": file_extension,
            "is_supported": f".{file_extension}" in ['.pdf', '.docx', '.doc', '.xlsx', '.xls']
        }